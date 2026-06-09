#!/usr/bin/env python3
"""Build LEVEL3_REPORT.docx — Level 3 focused progress report."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
NAVY = RGBColor(0x1F, 0x49, 0x7B)
doc = Document()
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)


def hdr_shade(cell, color='1F497B'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255); hdr_shade(c)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(10)
    doc.add_paragraph()
    return t


def h1(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = NAVY


def h2(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = True; r.font.size = Pt(12.5)
    r.font.color.rgb = NAVY


def body(txt):
    doc.add_paragraph(txt)


# ── Title ──
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('Level 3 Progress Report'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Resilient RTK for Disaster Rescue'); r.italic = True; r.font.size = Pt(13)
mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mp.add_run('UAV Emergency Rescue BDS — RTK Positioning Module  |  Tevin Wills  |  2026-06-04').font.size = Pt(10)
doc.add_paragraph()

h1('1.  Summary')
body('Level 3 tests whether the RTK positioning system can support an autonomous rescue '
     'mission in a disaster zone where GNSS is progressively and simultaneously degraded. '
     'We designed a single compound disaster scenario — not four isolated failures — flew '
     'it on a real PX4/Gazebo drone mission, and analysed the result against engineering '
     'thresholds for approach navigation and precision landing.')
body('Headline result: RTK held mean positioning error to 0.968 m under sustained '
     'multi-factor degradation — about 5x better than the total-failure case (4.776 m), '
     'and about 20x the ideal baseline (0.048 m). The system PASSED the approach threshold '
     '(92.4% of samples within 2.0 m, target >= 90%) but FELL SHORT of the precision-landing '
     'threshold (31.5% within 0.3 m, target >= 80%).')

h1('2.  What was done, and why')
add_table(['What', 'Why'], [
    ['Designed a 5-phase compound degradation profile', 'Real disaster zones degrade on several axes at once; isolated single-fault tests do not represent that'],
    ['Modelled 3 co-occurring degradations: noise 1.5->2.75 m, quality 0.95->0.60, dropouts up to 37.5 s / 90 s', 'To stress the system the way a collapsing-structure search would'],
    ['Added a dynamic uncertainty output and a mission viability state signal', 'A path planner needs an honest, time-varying confidence signal, not a fixed accuracy claim'],
    ['Flew the scenario as a real 5-waypoint QGC mission (390 x 789 m)', 'To validate against true vehicle flight, not a stationary hover'],
    ['Ran a total-failure control and reused the Level 2 ideal run as baseline', 'To bracket the disaster result between best and worst case'],
    ['Produced 7 analysis figures and cross-validated against the PX4 ULog', 'To make the result presentation-ready and provably real'],
])

h1('3.  Method (scenario at a glance)')
add_table(['Phase', 'Time (s)', 'GNSS noise', 'Quality', 'Dropouts'], [
    ['Departure', '0-120', '1.5 m', '0.95', 'none'],
    ['Approach', '120-240', '1.5 -> 2.5 m', '0.95 -> 0.60', '5 s / 60 s'],
    ['Search Zone', '240-600', '2.75 m', '0.60', '37.5 s / 90 s'],
    ['Landing', '600-660', '2.5 m', '0.60', 'none'],
    ['Exit', '660-780', '2.5 -> 1.5 m', '0.65 -> 0.90', 'none'],
])
body('Viability states: LANDING_VIABLE <= 0.3 m  ·  APPROACH_VIABLE <= 2.0 m  ·  '
     'DEGRADED <= 4.0 m  ·  INSUFFICIENT > 4.0 m.')

h1('4.  What was successfully confirmed')
for t in [
    'RTK degrades gracefully, not catastrophically — mean error 0.968 m (compound) vs 4.776 m (no corrections), ~5x more accurate even under disaster conditions. (Fig 2, Fig 5)',
    'Approach navigation is viable — 92.4% of approach-phase samples within 2.0 m, beating the >= 90% target. (Fig 6)',
    'The uncertainty signal is honest — reported rtk_std rises and falls with actual error (0.03 -> 4.6 m), unlike the flat Level 2 fixed value. (Fig 3)',
    'The viability state machine behaves correctly — downgrades during dropouts, recovers to LANDING_VIABLE as corrections return. (Fig 4)',
    'The data is real and self-consistent — the PX4 ULog (409 MB) and analysis CSV trace the same 5-waypoint trajectory. (Fig 7)',
]:
    doc.add_paragraph(t, style='List Number')

h1('5.  Honest limitations')
for t in [
    'Precision landing not met: only 31.5% of landing samples within 0.3 m (target 80%) — reflects RTK Float, not Fixed, during degradation. The key open engineering gap.',
    'Run altitude mismatch: the compound run cruised at 50 m throughout while the total-failure run descended to ~20 m, so the runs are not strictly identical flights. Per-run statistics are unaffected.',
    'Phase windows are time-reconstructed and lead the actual degradation onset by a few tens of seconds; Exit slightly overruns the designed 780 s.',
    'Baseline is a different (shorter) Level 2 mission — used as an accuracy-only reference, noted as such.',
]:
    doc.add_paragraph(t, style='List Number')

h1('6.  Next task')
for t in [
    'Close the landing gap — implement an RTK re-initialisation / Fixed-lock recovery protocol on entering the landing phase, or extend the landing window for re-convergence.',
    'Re-fly both Level 3 runs with one identical mission (ideally including the WP4 20 m descent) so the runs are strictly comparable.',
    'Phase 3 — RTKLIB validation: post-process real BeiDou RINEX data through RTKLIB and compare measured SPP / Float / Fixed accuracy against the simulation parameters.',
    'Push Level 3 work to the team GitHub once the above are scheduled.',
]:
    doc.add_paragraph(t, style='List Number')

h1('7.  Figure index')
add_table(['Fig', 'File', 'Purpose'], [
    ['1', 'l3_compound_scenario_profile.png', 'The imposed 5-phase degradation model'],
    ['2', 'l3_error_over_time.png', 'RTK vs raw error over the mission (log)'],
    ['3', 'l3_uncertainty_over_time.png', 'Dynamic vs fixed uncertainty (log)'],
    ['4', 'l3_viability_timeline.png', 'Mission viability state + error (log)'],
    ['5', 'l3_three_run_comparison.png', 'Baseline vs compound vs total failure'],
    ['6', 'l3_mission_phase_accuracy.png', 'Accuracy & threshold pass/fail per phase'],
    ['7', 'l3_ulog_crossval.png', 'ULog cross-validation (trajectory/error/altitude)'],
])

out = os.path.join(HERE, 'LEVEL3_REPORT.docx')
doc.save(out)
print('saved', out)
