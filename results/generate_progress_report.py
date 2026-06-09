"""
Generates the RTK Positioning Module Progress Report as a Word (.docx) document.
Run from any directory — all paths are absolute.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ────────────────────────────────────────────────────────────────────
BASE       = '/home/tevin_wills2/uav-emergency-rescue-bds'
IMG_L1     = os.path.join(BASE, 'results', 'graphs', 'rtk_positioning', 'level1')
IMG_L2     = os.path.join(BASE, 'results', 'graphs', 'rtk_positioning', 'level2')
IMG_L3     = os.path.join(BASE, 'results', 'graphs', 'rtk_positioning', 'level3')
OUT_PATH   = '/home/tevin_wills2/uav-emergency-rescue-bds/results/RTK_Positioning_Progress_Report.docx'

doc = Document()

# ── Page margins (2.5 cm all sides) ─────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def body(text, bold=False, italic=False, size=11, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, color=(31, 73, 125))
    p.paragraph_format.keep_with_next = True
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F497B')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_font(r, size=11.5, bold=True, color=(54, 96, 146))
    p.paragraph_format.keep_with_next = True
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    set_font(r, size=11, bold=True, italic=True, color=(79, 129, 189))
    p.paragraph_format.keep_with_next = True
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    r = p.add_run(text)
    set_font(r, size=11)
    return p

def add_figure(img_path, caption_num, caption_text, width=Inches(5.8)):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after  = Pt(2)
    run = p_img.add_run()
    run.add_picture(img_path, width=width)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after  = Pt(12)
    p_cap.paragraph_format.space_before = Pt(0)
    r_label = p_cap.add_run(f'Figure {caption_num}: ')
    set_font(r_label, size=9.5, bold=True, italic=True)
    r_text = p_cap.add_run(caption_text)
    set_font(r_text, size=9.5, italic=True)

def add_table_row(table, cells_data, header=False, bg_color=None):
    row = table.add_row()
    for i, (cell_text, width) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cell_text)
        set_font(r, size=10, bold=header)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if bg_color:
            tcPr = cell._tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)
    return row

def page_break():
    doc.add_page_break()

def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(pts)
    p.paragraph_format.space_before = Pt(0)

def inline_bold(para_obj, bold_text, normal_text=''):
    r1 = para_obj.add_run(bold_text)
    set_font(r1, bold=True, size=11)
    if normal_text:
        r2 = para_obj.add_run(normal_text)
        set_font(r2, size=11)

def make_header_row(table, headers):
    hdr_row = table.rows[0]
    for i, txt in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        set_font(r, size=10, bold=True)
        r.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497B')
        tcPr.append(shd)

def add_data_rows(table, rows_data, first_col_left=True):
    alt = False
    for rd in rows_data:
        row = table.add_row()
        fill = 'D6E4F0' if alt else 'FFFFFF'
        alt  = not alt
        for i, txt in enumerate(rd):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            set_font(r, size=10)
            tcPr = cell._tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

def table_caption(num_str, caption_text):
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after  = Pt(10)
    r1 = p_cap.add_run(f'Table {num_str}: ')
    set_font(r1, size=9.5, bold=True, italic=True)
    r2 = p_cap.add_run(caption_text)
    set_font(r2, size=9.5, italic=True)

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
spacer(60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('UAV Emergency Rescue System')
set_font(r, size=22, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('RTK Positioning Module')
set_font(r, size=18, bold=True, color=(54, 96, 146))

spacer(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Progress Report')
set_font(r, size=15, bold=False, color=(89, 89, 89))

spacer(40)

for label, value in [
    ('Student', 'Tevin Wills  (Student 1 — RTK Positioning)'),
    ('Module', 'RTK-Based Positioning Subsystem'),
    ('Simulation Stack', 'ROS 2 Jazzy · PX4 SITL · Gazebo · QGroundControl'),
    ('Date', '4 June 2026'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f'{label}:  ')
    set_font(r1, size=11, bold=True, color=(54, 96, 146))
    r2 = p.add_run(value)
    set_font(r2, size=11)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════
heading1('Table of Contents')

toc_items = [
    ('1.', 'Executive Summary', '3'),
    ('2.', 'Project Background and Module Scope', '3'),
    ('3.', 'System Architecture', '4'),
    ('4.', 'Level 1 — Standalone Simulation', '4'),
    ('5.', 'Level 2 — PX4/Gazebo Integration', '8'),
    ('6.', 'Level 3 — Resilient RTK for Disaster Rescue', '14'),
    ('7.', 'Challenges and Limitations — Simulation vs. Real RTK Hardware', '21'),
    ('8.', 'Current Status', '23'),
    ('9.', 'Next Steps', '24'),
]

for num, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(5.5), 1)  # right-align at 5.5"
    r1 = p.add_run(f'{num}  {title}')
    set_font(r1, size=10.5)
    r2 = p.add_run(f'\t{pg}')
    set_font(r2, size=10.5)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
heading1('1.   Executive Summary')

body(
    'This report presents the progress made on the RTK (Real-Time Kinematic) Positioning module '
    'of the UAV Emergency Rescue System across three simulation levels. Level 1 validated the RTK '
    'algorithm in a standalone ROS 2 environment, achieving a mean positioning accuracy of 0.101 m '
    'compared to a raw GNSS baseline of 2.394 m — a 95.8% improvement. Level 2 extended this to a '
    'full PX4/Gazebo simulation with an autonomous QGroundControl waypoint mission covering a '
    '208 × 214 m search area, achieving 0.048 m mean error during RTK Fixed lock — a 96.7% '
    'improvement and an 18.8-fold improvement over the autopilot\'s own GPS accuracy estimate. '
    'Level 3 introduced a compound disaster scenario simulating five phases of signal degradation '
    'across a 5-waypoint rescue mission covering a 390 × 789 m area. The system met the '
    'approach navigation threshold (92.4% of approach-phase samples within 2.0 m, against a '
    '≥ 90% target) but fell short of the precision landing threshold (31.5% within 0.3 m, '
    'against a ≥ 80% target) under peak search-zone interference — a genuine finding about '
    'system limitations under sustained degraded conditions. All three levels have been logged, '
    'analysed, and documented with publication-quality figures.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 2. PROJECT BACKGROUND
# ═══════════════════════════════════════════════════════════════════════════
heading1('2.   Project Background and Module Scope')

body(
    'The UAV Emergency Rescue System is a multi-module team project aimed at developing an '
    'autonomous UAV platform capable of locating and responding to distress signals in '
    'environments where ground-based rescue teams face access constraints. The system integrates '
    'several subsystems including computer vision, mission planning, communication, and positioning. '
    'This report covers the RTK Positioning module exclusively.'
)

body(
    'The objective of the RTK Positioning module is to provide accurate, reliable, and '
    'real-time position estimates to the UAV flight controller, significantly exceeding the '
    'accuracy of standard GNSS. In a rescue scenario, the difference between a 2.4 m positioning '
    'error and a 0.05 m error determines whether the UAV can precisely locate a survivor, '
    'deliver a payload to a specific point, or navigate safely in a constrained space. The module '
    'achieves this by combining raw GNSS measurements with RTCM differential corrections broadcast '
    'from a fixed ground base station.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
heading1('3.   System Architecture')

body(
    'The RTK Positioning module is implemented as a set of ROS 2 nodes communicating over '
    'CycloneDDS. The following nodes form the core of the system:'
)

bullet('Base Station Node — simulates a fixed-position GNSS reference receiver, publishing RTCM differential corrections over the /rtcm/corrections topic.')
bullet('RTCM Correction Simulator Node (Level 2 and 3) — models a realistic radio-link correction feed including correction age tracking, deliberate outage events, and in Level 3, a phased degradation model with variable quality and periodic dropouts.')
bullet('RTK Positioning Node — the primary processing node. It fuses raw GNSS measurements with RTCM corrections to produce RTK-corrected position estimates, managing fix-state transitions (GNSS Only → RTK Float → RTK Fixed → Correction Lost).')
bullet('PX4 Pose Adapter Node (Level 2 and 3) — bridges the PX4 SITL flight controller to the ROS 2 graph via MicroXRCEAgent, providing ground truth pose data from the Gazebo physics engine.')
bullet('Logger Node — subscribes to all positioning, ground truth, and mission viability topics, writing timestamped CSV log files at 10 Hz for post-flight analysis.')

body(
    'The system runs on ROS 2 Jazzy under Ubuntu 24.04 (Noble) on WSL2. In Levels 2 and 3, '
    'PX4 SITL and Gazebo run on the same host, with QGroundControl (QGC) connected via UDP for '
    'mission upload and monitoring.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. LEVEL 1
# ═══════════════════════════════════════════════════════════════════════════
heading1('4.   Level 1 — Standalone Simulation')

heading2('4.1   Experimental Setup')

body(
    'Level 1 validates the RTK algorithm independently of any flight controller, using a '
    'simulated UAV node to generate a deterministic ground truth trajectory. The UAV follows '
    'a 50 × 50 m square path at a fixed altitude of 30 m above ground level (AGL). The '
    'simulation runs at 10 Hz for a total duration of 271.7 seconds, producing 2,718 logged '
    'samples. The data file used for all analysis is rtk_level1_20260521_020841.csv.'
)

body(
    'Raw GNSS measurements are generated by adding Gaussian noise (σ = 1.5 m) to the ground '
    'truth position. RTCM corrections from the base station node progressively improve the '
    'position estimate as the fix state transitions from GNSS Only to RTK Float to RTK Fixed. '
    'All coordinates are expressed in East-North-Up (ENU) metres relative to the base station '
    'at latitude 39.981°N, longitude 116.344°E.'
)

heading2('4.2   RTK Status Progression')

body(
    'The system progresses through three fix states during the Level 1 run. Table 4.1 summarises '
    'the duration, sample count, and positioning accuracy achieved in each state.'
)

# Table 4.1
tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.paragraph_format = None

hdr_data = [
    ('Phase', 1.2), ('Time Window', 1.2), ('Samples (%)', 1.2),
    ('Raw GNSS Error', 1.4), ('RTK-Corrected Error', 1.6),
]
hdr_row = tbl.rows[0]
for i, (txt, _) in enumerate(hdr_data):
    cell = hdr_row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    set_font(r, size=10, bold=True)
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497B')
    tcPr.append(shd)
    r.font.color.rgb = RGBColor(255, 255, 255)

rows_data = [
    ['GNSS Only',  '0 – 5 s',    '51  (1.9%)',   '2.121 m', '2.192 m'],
    ['RTK Float',  '5 – 15 s',   '100  (3.7%)',  '2.524 m', '0.389 m'],
    ['RTK Fixed',  '≥ 15 s',     '2,567  (94.4%)','2.394 m', '0.048 m'],
]
add_data_rows(tbl, rows_data)
table_caption('4.1', 'RTK fix state distribution and mean positioning error — Level 1 standalone simulation.')

heading2('4.3   Positioning Accuracy Results')

body(
    'Across the full 271.7-second run, the RTK positioning system achieved a mean 3D error of '
    '0.101 m against simulated ground truth, compared to a raw GNSS mean error of 2.394 m. '
    'This represents a 95.8% reduction in positioning error. During the steady RTK Fixed phase '
    '(94.4% of the run), the mean RTK-corrected error was 0.048 m with a standard deviation of '
    '0.021 m and a 95th-percentile error of 0.085 m.'
)

body(
    'Figure 4.1 shows the full error time series over the 271.7-second run. The raw GNSS error '
    '(red) varies between 0.5 m and 6.2 m, reflecting the Gaussian noise model. The RTK-corrected '
    'error (blue) drops sharply as the fix state progresses, settling near zero once RTK Fixed is '
    'achieved at t = 15 s. The three background shading regions identify the GNSS Only (amber), '
    'RTK Float (yellow), and RTK Fixed (green) phases.'
)

add_figure(
    os.path.join(IMG_L1, 'l1_error_over_time.png'),
    '4.1',
    'RTK positioning error over time — Level 1 standalone simulation. Red: raw GNSS error '
    '(3 s rolling mean). Blue: RTK-corrected error (3 s rolling mean). Dotted lines indicate '
    'overall means. Dashed lines indicate theoretical accuracy specifications per fix state.'
)

body(
    'Figure 4.2 focuses on the first 60 seconds of the run, capturing the complete convergence '
    'sequence from GNSS Only through to RTK Fixed. The lower status strip confirms the timing of '
    'each fix-state transition. The raw GNSS error remains largely unchanged during convergence, '
    'while the RTK-corrected error reduces progressively as more correction data is accumulated.'
)

add_figure(
    os.path.join(IMG_L1, 'l1_rtk_convergence.png'),
    '4.2',
    'RTK fix convergence during the first 60 seconds — Level 1 standalone simulation. '
    'Vertical dotted lines mark the GNSS Only → RTK Float transition at t = 5 s and the '
    'RTK Float → RTK Fixed transition at t = 15 s. The lower strip shows the fix status timeline.'
)

body(
    'Figure 4.3 presents the error distribution for both raw GNSS and RTK-corrected measurements, '
    'using only RTK Fixed samples to represent steady-state performance. The raw GNSS distribution '
    '(left panel) is broad, with a KDE peak near 2.41 m and a 95th-percentile error of 4.18 m. '
    'The RTK-corrected distribution (right panel) is narrow and tightly concentrated near 0.048 m, '
    'with a 95th-percentile error of 0.085 m — confirming consistent sub-decimetre performance '
    'across the steady-state phase.'
)

add_figure(
    os.path.join(IMG_L1, 'l1_error_distribution.png'),
    '4.3',
    'Positioning error distribution during the RTK Fixed period — Level 1 standalone simulation. '
    'Left: raw GNSS error (0–8 m scale). Right: RTK-corrected error (0–0.5 m scale). '
    'Dashed line: mean. Dotted line: median. Dash-dot line: 95th percentile.'
)

body(
    'Figure 4.4 shows the top-down 2D trajectory in the ENU frame. The ground truth path (dark '
    'line) traces the 50 × 50 m square perimeter. The raw GNSS scatter (red) is visibly dispersed '
    'up to several metres around the true path, while the RTK-corrected scatter (blue) is tightly '
    'clustered along the ground truth. The four corner markers (NE, NW, SE, SW) confirm the '
    'planned path geometry was followed correctly throughout the run.'
)

add_figure(
    os.path.join(IMG_L1, 'l1_trajectory.png'),
    '4.4',
    '2D top-down UAV trajectory in the ENU frame — Level 1 standalone simulation. '
    'Dark line: ground truth path. Blue scatter: RTK-corrected positions. '
    'Red scatter: raw GNSS positions. Purple circle: start/end point.'
)

body(
    'Figure 4.5 provides a consolidated performance summary. The left panel compares mean raw '
    'GNSS and RTK-corrected errors side by side for each fix phase. The centre panel shows the '
    'distribution of time spent in each fix state. The right panel shows RTK-corrected error per '
    'phase with standard deviation error bars, and confirms the 95.8% overall accuracy improvement.'
)

add_figure(
    os.path.join(IMG_L1, 'l1_accuracy_summary.png'),
    '4.5',
    'RTK positioning performance summary — Level 1 standalone simulation. '
    'Left: grouped mean error by fix phase (raw GNSS vs RTK-corrected). '
    'Centre: fix state time distribution. Right: RTK-corrected error per fix phase with ±1σ error bars.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. LEVEL 2
# ═══════════════════════════════════════════════════════════════════════════
heading1('5.   Level 2 — PX4/Gazebo Integration')

heading2('5.1   Experimental Setup')

body(
    'Level 2 extends the standalone simulation by integrating the RTK positioning module with a '
    'full PX4 Software-In-The-Loop (SITL) flight stack running inside Gazebo. The ROS 2 – PX4 '
    'bridge is provided by MicroXRCEAgent over DDS. QGroundControl (QGC) was used to upload and '
    'execute an autonomous waypoint mission consisting of eight waypoints covering a 208 × 214 m '
    'search area at an altitude of 50 m AGL. The mission ran for 166 seconds of active flight. '
    'Six ROS 2 nodes ran concurrently: base station, RTK positioning, RTCM correction simulator, '
    'PX4 pose adapter, logger, and the px4_pose_adapter bridge node.'
)

body(
    'The simulation log captured 10,282 samples at 10 Hz over 1,028 seconds total, including '
    'pre-flight initialisation, the full mission, and post-landing data. The data file used for '
    'analysis is rtk_level2_20260521_022231.csv. The QGC flight log '
    '(qgc_mission_20260521.ulg) was used for independent cross-validation.'
)

heading2('5.2   RTK Status Progression')

body(
    'The system progressed through four distinct fix states during the Level 2 run, including '
    'a deliberate correction outage event. Table 5.1 summarises each phase.'
)

# Table 5.1
tbl2 = doc.add_table(rows=1, cols=5)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, txt in enumerate(['Phase', 'Time Window', 'Samples (%)', 'Raw GNSS Error', 'RTK-Corrected Error']):
    cell = tbl2.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    set_font(r, size=10, bold=True)
    r.font.color.rgb = RGBColor(255, 255, 255)
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497B')
    tcPr.append(shd)

rows2 = [
    ['GNSS Only',        '0 – 5 s',       '50  (0.5%)',    '2.121 m', '2.192 m'],
    ['RTK Float',        '5 – 15 s',      '100  (1.0%)',   '2.524 m', '0.389 m'],
    ['RTK Fixed',        '15–45 s, ≥50 s','10,082  (98.1%)','2.412 m', '0.048 m'],
    ['Correction Lost',  '45 – 50 s',     '50  (0.5%)',    '2.598 m', '3.774 m'],
]
add_data_rows(tbl2, rows2)
table_caption('5.1', 'RTK fix state distribution and mean positioning error — Level 2 PX4/Gazebo simulation.')

body(
    'The Correction Lost phase at t = 45–50 s is a notable event. RTCM corrections were '
    'temporarily unavailable for five seconds, causing the RTK error to spike to a peak of '
    '7.46 m before the system autonomously re-acquired RTK Fixed lock at t = 50 s. This '
    'demonstrates the system\'s resilience under a correction outage — a realistic scenario '
    'in field deployments where base-station radio links can be intermittently disrupted.'
)

heading2('5.3   Positioning Accuracy Results')

body(
    'Across the full 1,028-second run, the RTK positioning system achieved a mean 3D error of '
    '0.080 m against the Gazebo ground truth, compared to a raw GNSS mean error of 2.413 m. '
    'This corresponds to an accuracy improvement of 96.7%. During the steady RTK Fixed phase '
    '(98.1% of the run), the mean error was 0.048 m with a standard deviation of 0.020 m and '
    'a 95th-percentile error of 0.083 m.'
)

body(
    'Figure 5.1 shows the full error time series over the 1,028-second run. The Correction Lost '
    'spike at t = 48.6 s is annotated. Outside this brief event, the RTK-corrected error '
    'remains consistently below 0.15 m throughout the entire mission.'
)

add_figure(
    os.path.join(IMG_L2, 'l2_error_over_time.png'),
    '5.1',
    'RTK positioning error over time — Level 2 PX4/Gazebo simulation. Red: raw GNSS error '
    '(3 s rolling mean). Blue: RTK-corrected error (3 s rolling mean). The annotated spike at '
    't ≈ 49 s corresponds to the Correction Lost event. Background shading identifies the four '
    'fix state phases.'
)

body(
    'Figure 5.2 focuses on the first 60 seconds, capturing all four phase transitions including '
    'the brief Correction Lost period at t = 45–50 s. The status strip at the bottom confirms '
    'the exact timing of each transition. The RTK-corrected error is seen to spike and recover '
    'sharply within the five-second window.'
)

add_figure(
    os.path.join(IMG_L2, 'l2_rtk_convergence.png'),
    '5.2',
    'RTK fix convergence during the first 60 seconds — Level 2 PX4/Gazebo simulation. '
    'All four fix-state transitions are marked. The bottom strip shows the complete status '
    'timeline including the five-second Correction Lost gap at t = 45–50 s.'
)

body(
    'Figure 5.3 shows the error distribution for RTK Fixed samples only, confirming sub-decimetre '
    'steady-state performance. The RTK-corrected distribution (right panel) peaks at 0.047 m '
    'with negligible probability mass beyond 0.15 m. The raw GNSS distribution (left panel) '
    'is broad and centred near 2.41 m, as expected for uncorrected pseudorange measurements.'
)

add_figure(
    os.path.join(IMG_L2, 'l2_error_distribution.png'),
    '5.3',
    'Positioning error distribution during the RTK Fixed period — Level 2 PX4/Gazebo simulation. '
    'Left: raw GNSS error (0–8 m scale). Right: RTK-corrected error (0–0.5 m scale).'
)

body(
    'Figure 5.4 shows the 2D top-down trajectory of the autonomous QGC mission. The mission '
    'covers a significantly larger area than Level 1, spanning 208 m in the East direction and '
    '214 m in the North direction. The RTK-corrected scatter (blue) remains tightly clustered '
    'around the ground truth path (dark line) throughout the full mission extent.'
)

add_figure(
    os.path.join(IMG_L2, 'l2_trajectory.png'),
    '5.4',
    '2D top-down UAV flight trajectory — Level 2 PX4/Gazebo simulation, autonomous QGC mission. '
    'Dark line: ground truth path. Blue scatter: RTK-corrected positions. '
    'Red scatter: raw GNSS positions. Purple circle: home/takeoff point.'
)

body(
    'Figure 5.5 provides the consolidated performance summary for Level 2, including the '
    'four-phase fix state distribution and RTK-corrected error per phase. The Correction Lost '
    'bar clearly shows the degradation to 3.77 m during the outage, contrasting with the '
    '0.048 m achieved during RTK Fixed.'
)

add_figure(
    os.path.join(IMG_L2, 'l2_accuracy_summary.png'),
    '5.5',
    'RTK positioning performance summary — Level 2 PX4/Gazebo simulation. '
    'Left: grouped mean error by fix phase. Centre: four-phase fix state distribution. '
    'Right: RTK-corrected error per phase with ±1σ error bars and overall improvement annotation.'
)

heading2('5.4   QGC ULog Cross-Validation')

body(
    'The QGC flight log (qgc_mission_20260521.ulg) was parsed using the pyulog library to '
    'independently verify the mission trajectory and to benchmark the RTK system against '
    'PX4\'s own GPS accuracy estimate. Three findings are reported.'
)

p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.left_indent  = Inches(0.25)
r = p.add_run('Trajectory confirmation. ')
set_font(r, bold=True, size=11)
r2 = p.add_run(
    'The ground truth path recorded in the ULog matches the Level 2 CSV ground truth geometry '
    '(X range −63.8 to +144.8 m, Y range −81.1 to +133.6 m), confirming that both data sources '
    'describe the same Gazebo simulation flight.'
)
set_font(r2, size=11)

p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.left_indent  = Inches(0.25)
r = p.add_run('PX4 GPS accuracy (EPH). ')
set_font(r, bold=True, size=11)
r2 = p.add_run(
    'The PX4 autopilot reported a horizontal position error estimate (EPH) of 0.90 m '
    'throughout the mission. EPH is the accuracy figure the autopilot uses for its own '
    'navigation decisions and is a standard metric for comparing positioning system performance.'
)
set_font(r2, size=11)

p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(8)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.left_indent  = Inches(0.25)
r = p.add_run('RTK advantage over autopilot GPS. ')
set_font(r, bold=True, size=11)
r2 = p.add_run(
    'The RTK system achieved a measured error of 0.048 m during RTK Fixed, which is an '
    '18.8-fold improvement over the autopilot\'s EPH of 0.90 m. In an emergency rescue scenario, '
    'this translates to the difference between locating a survivor within a metre versus '
    'within five centimetres.'
)
set_font(r2, size=11)

body(
    'Figure 5.6 presents the full cross-validation result across three panels: the ULog ground '
    'truth trajectory with PX4 GPS positions and mission waypoints; the three-way accuracy '
    'comparison bar chart (Raw GNSS 2.413 m, PX4 GPS EPH 0.900 m, RTK Fixed 0.048 m); and '
    'the UAV altitude profile from the ULog, confirming the complete mission was executed '
    'correctly with takeoff, cruise at 50 m, and landing.'
)

add_figure(
    os.path.join(IMG_L2, 'l2_qgc_crossval.png'),
    '5.6',
    'QGC ULog cross-validation — Level 2. Left: ULog ground truth trajectory with PX4 GPS '
    'scatter and mission waypoints (triangles). Centre: three-way accuracy comparison '
    '(raw GNSS · PX4 GPS EPH · RTK Fixed). Right: UAV altitude profile confirming autonomous '
    'mission execution with takeoff and landing events marked.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. LEVEL 3
# ═══════════════════════════════════════════════════════════════════════════
heading1('6.   Level 3 — Resilient RTK for Disaster Rescue')

heading2('6.1   Experimental Setup and Mission Design')

body(
    'Level 3 tests the RTK positioning system under compound disaster conditions: simultaneous '
    'GNSS degradation and RTCM correction link disruption across a realistic rescue mission. '
    'The evaluation uses a three-run comparison design. Run 1 (baseline) is the Level 2 CSV '
    'result under ideal conditions. Run 2 is the compound disaster scenario — a five-phase '
    'signal degradation profile applied to a full autonomous mission. Run 3 is a total failure '
    'scenario — the worst credible case in which no RTCM corrections are ever received and GNSS '
    'noise is fixed at 3.0 m throughout.'
)

body(
    'The mission is a five-waypoint autonomous QGC flight covering 390 m East × 789 m North at '
    '50 m AGL, flown at 5 m/s cruise speed. This is significantly larger than the Level 2 mission '
    '(208 × 214 m) and represents a realistic single-operator UAV rescue search area. The mission '
    'was flown in PX4 SITL with Gazebo, with the same ROS 2 node stack as Level 2 plus the '
    'Level 3 RTCM correction simulator that implements the five-phase degradation model.'
)

body(
    'Two engineering thresholds were defined prior to data collection to evaluate operational '
    'mission viability:'
)
bullet('Approach navigation threshold: ≥ 90% of approach-phase samples within 2.0 m RTK error.')
bullet('Precision landing threshold: ≥ 80% of landing-phase samples within 0.3 m RTK error.')

body(
    'These thresholds are derived from operational UAV rescue requirements: approach accuracy '
    'determines safe area navigation, while precision landing accuracy determines whether the '
    'UAV can safely deliver a payload or initiate a survivor approach to within arm\'s reach.'
)

heading2('6.2   Five-Phase Signal Degradation Model')

body(
    'The compound disaster scenario divides the mission into five phases, each with a distinct '
    'GNSS noise level and RTCM correction quality. Table 6.1 summarises the phase parameters. '
    'The search zone phase (240–600 s) includes periodic correction dropouts at 60-second '
    'intervals to simulate intermittent base-station radio link failure in a disaster area. '
    'Figure 6.1 shows the scenario profile as recorded in the compound disaster data run.'
)

# Table 6.1 — Phase profile
tbl_l3phase = doc.add_table(rows=1, cols=5)
tbl_l3phase.style = 'Table Grid'
tbl_l3phase.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, txt in enumerate(['Phase', 'Mission Time', 'GNSS Noise (σ)', 'Correction Quality', 'Notes']):
    cell = tbl_l3phase.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    set_font(r, size=10, bold=True)
    r.font.color.rgb = RGBColor(255, 255, 255)
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497B')
    tcPr.append(shd)

phase_rows = [
    ['Departure',   '0 – 120 s',   '1.50 m', '0.95', 'Near-baseline conditions'],
    ['Approach',    '120 – 240 s', '2.00 m', '0.80', 'Moderate degradation'],
    ['Search Zone', '240 – 600 s', '2.75 m', '0.60', 'Worst phase — periodic dropouts'],
    ['Landing',     '600 – 660 s', '1.50 m', '0.95', 'Temporarily restored'],
    ['Exit',        '660+ s',      '1.50 m', '0.95', 'Return to base conditions'],
]
add_data_rows(tbl_l3phase, phase_rows)
table_caption('6.1', 'Level 3 compound disaster five-phase signal degradation profile. '
              'GNSS noise and correction quality vary by mission phase to simulate realistic '
              'disaster-area RF conditions.')

add_figure(
    os.path.join(IMG_L3, 'l3_compound_scenario_profile.png'),
    '6.1',
    'Level 3 compound disaster scenario profile over mission elapsed time. '
    'Top panel: GNSS noise standard deviation (m). Middle panel: RTCM correction quality (0–1). '
    'Bottom strip: mission phase colour coding. The five-phase structure and search-zone '
    'degradation peak are clearly visible.'
)

heading2('6.3   Error and Uncertainty Time Series')

body(
    'Figure 6.2 shows the RTK positioning error over the full compound disaster mission. '
    'The error trace follows the five-phase structure directly: near-baseline during Departure, '
    'rising through the Approach phase, reaching peak values in the Search Zone (mean 1.86 m, '
    'max 9.79 m), and recovering at Landing. The correction dropout events in the Search Zone '
    'produce the characteristic error spikes visible as the RTK status transitions to '
    'Correction Lost. The mean error across the full compound disaster run is 0.968 m — '
    'twenty times higher than the baseline 0.048 m — demonstrating the impact of sustained '
    'multi-factor degradation.'
)

add_figure(
    os.path.join(IMG_L3, 'l3_error_over_time.png'),
    '6.2',
    'RTK positioning error over mission elapsed time — Level 3 compound disaster scenario. '
    'Rolling mean (3 s window) shown in blue. Background shading identifies mission phase. '
    'Error spikes in the Search Zone correspond to RTCM correction dropout events. '
    'Mean (dashed) and mission phase boundaries (vertical dotted) annotated.'
)

body(
    'Figure 6.3 shows the RTK position uncertainty (standard deviation of the RTK estimate) '
    'over the same mission. Uncertainty tracks the correction quality closely: low uncertainty '
    'during Departure and Landing phases (RTK Fixed conditions, σ ≈ 0.03–0.05 m) and elevated '
    'uncertainty during the Search Zone (σ ≈ 0.25–2.50 m as corrections degrade). The mission '
    'viability thresholds (0.30 m and 2.00 m) are overlaid, showing which viability states '
    'apply at each point in the mission.'
)

add_figure(
    os.path.join(IMG_L3, 'l3_uncertainty_over_time.png'),
    '6.3',
    'RTK position uncertainty (σ, metres) over mission elapsed time — Level 3 compound disaster. '
    'Horizontal dashed lines show mission viability thresholds: 0.30 m (LANDING_VIABLE) and '
    '2.00 m (APPROACH_VIABLE). Phases above 2.00 m are classified DEGRADED or INSUFFICIENT. '
    'Background shading identifies mission phase.'
)

heading2('6.4   Mission Viability Assessment')

body(
    'Figure 6.4 shows the mission viability timeline — the sequence of viability states '
    '(LANDING_VIABLE, APPROACH_VIABLE, DEGRADED, INSUFFICIENT) as a function of mission '
    'elapsed time. During the Departure and Landing phases, the system operates in '
    'LANDING_VIABLE state. During the Approach phase, the system is predominantly in '
    'APPROACH_VIABLE. In the Search Zone, the system transitions between APPROACH_VIABLE, '
    'DEGRADED, and INSUFFICIENT during correction dropout events.'
)

add_figure(
    os.path.join(IMG_L3, 'l3_viability_timeline.png'),
    '6.4',
    'Mission viability state timeline — Level 3 compound disaster scenario. '
    'Each horizontal band represents a mission viability state: LANDING_VIABLE (green), '
    'APPROACH_VIABLE (yellow-green), DEGRADED (orange), INSUFFICIENT (red). '
    'Vertical dotted lines mark phase transitions. The INSUFFICIENT state appears only '
    'during correction dropout events in the Search Zone.'
)

body(
    'The engineering threshold assessment yields the following results for the compound '
    'disaster scenario:'
)

p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.left_indent  = Inches(0.25)
r = p.add_run('Approach threshold: 92.4% within 2.0 m  ✔  PASS')
set_font(r, bold=True, size=11, color=(31, 120, 31))
r2 = p.add_run(
    '  —  The system exceeded the ≥ 90% target. During the Approach phase (120–240 s mission '
    'elapsed), moderate degradation (GNSS σ = 2.0 m, correction quality 0.80) produces errors '
    'predominantly in the APPROACH_VIABLE range. The 92.4% figure confirms the system '
    'remains operationally safe for area navigation despite elevated noise levels.'
)
set_font(r2, size=11)

p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(8)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.left_indent  = Inches(0.25)
r = p.add_run('Landing threshold: 31.5% within 0.3 m  ✘  FAIL')
set_font(r, bold=True, size=11, color=(160, 32, 32))
r2 = p.add_run(
    '  —  The system did not meet the ≥ 80% target for the landing phase in isolation. '
    'Although the Landing phase (600–660 s) nominally restores good correction quality, '
    'the RTK state machine requires re-convergence time after the prolonged Search Zone '
    'disruption. The 31.5% figure reflects RTK Float (not RTK Fixed) conditions during '
    'landing, where the position uncertainty remains above the 0.3 m precision landing '
    'threshold. This is a genuine system limitation: sustained search-zone interference '
    'prevents the rapid recovery needed for a precision landing immediately upon entering '
    'the landing phase.'
)
set_font(r2, size=11)

heading2('6.5   Three-Run Comparison')

body(
    'Figure 6.5 compares all three runs side by side using box plots, showing the full '
    'error distribution (median, interquartile range, outliers) for each run. Table 6.2 '
    'summarises the key statistics. The baseline (Level 2, ideal conditions) establishes '
    'the best-possible RTK performance. The compound disaster run demonstrates that the '
    'system remains broadly functional — albeit with a 20× increase in mean error — even '
    'under simultaneous GNSS and correction link degradation. The total failure run, with '
    'no corrections and elevated noise throughout, confirms the necessity of RTCM corrections: '
    'without them, mean error rises to 4.776 m — equivalent to raw GNSS performance under '
    'disaster conditions, and entirely incompatible with precision rescue navigation.'
)

# Table 6.2 — Three-run comparison
tbl_3run = doc.add_table(rows=1, cols=5)
tbl_3run.style = 'Table Grid'
tbl_3run.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, txt in enumerate(['Run', 'Samples', 'Mean RTK Error', '95th-Pct Error', 'Primary Status']):
    cell = tbl_3run.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    set_font(r, size=10, bold=True)
    r.font.color.rgb = RGBColor(255, 255, 255)
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497B')
    tcPr.append(shd)

runs_rows = [
    ['Run 1 — Baseline (L2)',        '10,282',  '0.048 m',   '0.083 m',  'RTK Fixed (98.1%)'],
    ['Run 2 — Compound Disaster',    '8,790',   '0.968 m',   '4.98 m',   'RTK Fixed / Float / Corr. Lost'],
    ['Run 3 — Total Failure',        '8,431',   '4.776 m',   '8.34 m',   'GNSS Only (99.4%)'],
]
add_data_rows(tbl_3run, runs_rows)
table_caption('6.2', 'Three-run comparison — Level 3 compound disaster scenario. '
              'Run 1 is the Level 2 baseline (ideal conditions). Run 2 is the compound disaster run. '
              'Run 3 is the total failure (no corrections, GNSS noise 3.0 m fixed throughout).')

add_figure(
    os.path.join(IMG_L3, 'l3_three_run_comparison.png'),
    '6.5',
    'Three-run RTK error comparison — box plots showing full error distribution for all three '
    'Level 3 runs. Box: interquartile range (IQR). Whiskers: 5th–95th percentile. Points beyond '
    'the whiskers are individual samples. The baseline (Run 1) shows near-zero median error. The compound disaster (Run 2) '
    'shows a substantially higher median with a long upper tail. The total failure (Run 3) '
    'shows a median around 4 m, representative of raw GNSS under disaster GNSS noise.'
)

heading2('6.6   Mission Phase Accuracy Analysis')

body(
    'Figure 6.6 disaggregates the compound disaster run by mission phase, showing the RTK error '
    'distribution for each of the five phases. This breakdown reveals which phase contributes '
    'most to overall error and confirms the five-phase structure visible in the time series. '
    'The Departure and Landing phases achieve near-baseline accuracy (median < 0.15 m), '
    'confirming that the system operates at RTK Fixed quality when correction quality is high. '
    'The Search Zone phase shows the highest spread, with error distributions extending well '
    'above 2.0 m during correction dropout windows. The Approach phase sits between the two '
    'extremes, consistent with the intermediate degradation parameters in the phase profile.'
)

add_figure(
    os.path.join(IMG_L3, 'l3_mission_phase_accuracy.png'),
    '6.6',
    'RTK error distribution by mission phase — Level 3 compound disaster scenario. '
    'Box plots show the error distribution within each of the five mission phases. '
    'Dashed horizontal lines show the mission viability thresholds: 0.30 m (precision landing) '
    'and 2.00 m (approach navigation). The Search Zone phase shows the widest distribution, '
    'reflecting periodic correction dropout events and peak GNSS noise.'
)

heading2('6.7   ULog Cross-Validation')

body(
    'The Level 3 compound disaster ULog (rtk_level3_compound_disaster_20260603.ulg, 409 MB) '
    'was parsed using pyulog to independently verify the Level 3 mission trajectory. The ULog '
    'records the vehicle\'s true Gazebo position in NED coordinates (x = North, y = East, '
    'z = Down). To align with the CSV ground truth convention (ENU: x = East, y = North), '
    'the mapping East = ULog.y and North = ULog.x was applied before comparison.'
)

body(
    'Figure 6.7 presents the cross-validation result. The left panel overlays the ULog NED '
    'trajectory (converted to ENU) against the CSV ground truth from the compound disaster '
    'data file, confirming geometric agreement: both sources show the same five-waypoint '
    'rectangular mission pattern spanning 390 m East × 789 m North. The right panel shows '
    'the UAV altitude profile from the ULog (-z axis), confirming cruise altitude at 50 m AGL '
    'throughout the mission with the expected takeoff and landing profiles. The trajectory '
    'agreement between ULog and CSV sources independently validates the integrity of the '
    'compound disaster data collection.'
)

add_figure(
    os.path.join(IMG_L3, 'l3_ulog_crossval.png'),
    '6.7',
    'ULog cross-validation — Level 3 compound disaster scenario. '
    'Left: ULog NED trajectory (converted to ENU, orange) overlaid on CSV ground truth (blue). '
    'Waypoint markers (triangles) from navigator_mission_item confirm the five-waypoint '
    'mission geometry. Right: UAV altitude profile from ULog (-z, NED), confirming 50 m AGL '
    'cruise and correct takeoff and landing profiles.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. CHALLENGES AND LIMITATIONS
# ═══════════════════════════════════════════════════════════════════════════
heading1('7.   Challenges and Limitations — Simulation vs. Real RTK Hardware')

body(
    'A core challenge of this phase of the project is that the RTK positioning module has been '
    'validated entirely within a software simulation environment. While the simulation demonstrates '
    'the correct functional behaviour of the system, several important physical phenomena present '
    'in real RTK deployments are either simplified or absent. Understanding these gaps is essential '
    'for interpreting the results and planning the transition to real hardware.'
)

heading2('7.1   Simplified GNSS Noise Model')

body(
    'The simulation models raw GNSS error as Gaussian noise with a fixed standard deviation '
    '(σ = 1.5 m during GNSS Only, reducing upon RTK convergence; up to 2.75 m in the Level 3 '
    'Search Zone phase). Real GNSS receivers are affected by multipath interference — signals '
    'reflecting off buildings, terrain, and vegetation before reaching the antenna. Multipath '
    'produces non-Gaussian, spatially correlated errors that cannot be cancelled by RTK '
    'corrections. In urban or forested environments, real raw GNSS errors can reach 5–15 m, '
    'far beyond the worst-case 2.75 m simulated here.'
)

heading2('7.2   Atmospheric Effects Not Modelled')

body(
    'Real RTK accuracy is influenced by ionospheric and tropospheric delays, which vary with '
    'time of day, season, solar activity, and weather conditions. RTK corrections from a base '
    'station partially cancel these effects for short baselines, but residual errors remain. '
    'Furthermore, real RTK accuracy degrades with baseline distance at approximately 1 part '
    'per million of separation — flying 789 m from the base station introduces roughly '
    '0.8 mm of additional baseline-dependent error. The simulation applies a phase-based '
    'noise model rather than a baseline-distance model, which is an approximation.'
)

heading2('7.3   Deterministic Convergence vs. Real Integer Ambiguity Resolution')

body(
    'In the simulation, the transition from RTK Float to RTK Fixed occurs deterministically '
    'at t = 15 s. In reality, RTK Fixed requires carrier-phase integer ambiguity resolution — '
    'a mathematically complex process dependent on satellite geometry (PDOP), baseline length, '
    'multipath conditions, and atmospheric stability. Real convergence can take anywhere from '
    '30 seconds to several minutes and can fail entirely in poor conditions. The Level 3 '
    'landing threshold failure (31.5% vs. ≥ 80% target) is partly attributable to this '
    'limitation: in reality, re-convergence after prolonged outage may take longer, making '
    'the threshold harder to meet still.'
)

heading2('7.4   Communication Link Idealisation')

body(
    'RTCM correction data in a real system is transmitted over a radio link, typically LoRa at '
    '433 or 915 MHz. This link is subject to range limitations, packet loss, channel fading, '
    'and interference from other RF sources on the UAV platform itself. The Level 3 correction '
    'dropout model (periodic, deterministic gaps in the Search Zone) is a simplification of '
    'the stochastic outage patterns observed in real field deployments. In real situations, '
    'outages are unpredictable in timing and duration.'
)

heading2('7.5   Hardware Effects Not Captured')

body('Several physical hardware factors absent from the simulation would affect real-world RTK performance:')

bullet('Antenna phase centre offset and variation — real patch antennas introduce position-dependent errors at low satellite elevation angles, which are not modelled in the simulation.')
bullet('Receiver noise floor — real GNSS receivers have hardware-specific noise characteristics that affect carrier-phase tracking quality and ultimately the achievable RTK accuracy.')
bullet('UAV vibration — mechanical vibration from rotors can cause antenna movement and degrade carrier-phase lock, particularly during aggressive manoeuvres or in windy conditions.')
bullet('Satellite geometry (PDOP) — the simulation uses a fixed noise model independent of satellite count or geometry. In reality, PDOP varies throughout the day and directly affects positioning accuracy.')

heading2('7.6   Implication for Next Steps')

body(
    'The simulation results establish a controlled performance baseline: 0.048 m mean error '
    'under ideal conditions (Level 2), rising to 0.968 m under compound disaster conditions '
    '(Level 3 Run 2), and 4.776 m under total loss of corrections (Level 3 Run 3). The '
    'planned Phase 3 RTKLIB validation against real BeiDou RINEX data from IGS MGEX stations '
    'will provide an independent calibration of the simulation noise model and establish '
    'whether the simulated degradation parameters (σ = 1.5–2.75 m) are representative of '
    'real disaster-area GNSS conditions.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. CURRENT STATUS
# ═══════════════════════════════════════════════════════════════════════════
heading1('8.   Current Status')

body('Table 8.1 summarises the completion status of all planned RTK positioning module milestones.')

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, txt in enumerate(['Milestone', 'Status', 'Notes']):
    cell = tbl3.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    set_font(r, size=10, bold=True)
    r.font.color.rgb = RGBColor(255, 255, 255)
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497B')
    tcPr.append(shd)

status_rows = [
    ('ROS 2 node architecture design',                    '✔  Complete',   'All 6 nodes implemented and tested'),
    ('Level 1 standalone simulation',                     '✔  Complete',   '95.8% accuracy improvement confirmed'),
    ('Level 1 data analysis and figures',                 '✔  Complete',   '5 publication-quality figures generated'),
    ('Level 2 PX4/Gazebo integration',                    '✔  Complete',   'All 6 nodes running with PX4 SITL'),
    ('Level 2 autonomous QGC mission',                    '✔  Complete',   '8-waypoint, 166 s mission executed'),
    ('Level 2 data analysis and figures',                 '✔  Complete',   '6 figures including ULog cross-validation'),
    ('Level 3 compound disaster scenario design',         '✔  Complete',   '5-phase degradation model implemented'),
    ('Level 3 QGC mission execution (compound + total failure)', '✔  Complete', '5-wp mission, 390 × 789 m, 2 valid runs'),
    ('Level 3 data analysis and figures',                 '✔  Complete',   '7 figures including ULog cross-validation'),
    ('Level 3 engineering threshold evaluation',          '✔  Complete',   'Approach ✔ 92.4%; Landing ✘ 31.5%'),
    ('GitHub push — Levels 1–3',                          '⏳  Pending',   'To be committed after report sign-off'),
    ('Phase 3 RTKLIB validation — plan',                  '⏳  Pending',   'BeiDou RINEX data from IGS MGEX stations'),
    ('Phase 3 RTKLIB validation — execution',             '○  Not started','Download RINEX, run rnx2rtkp, extract stats'),
    ('Integration with other team simulation modules',    '○  Not started','Navigation and mission planning modules'),
    ('Full system integration test in Gazebo',            '○  Not started','All team modules running simultaneously'),
]
alt = False
for milestone, stat, notes in status_rows:
    row = tbl3.add_row()
    fill = 'D6E4F0' if alt else 'FFFFFF'
    alt  = not alt
    color = '1F7A1F' if '✔' in stat else ('7A5200' if '⏳' in stat else '555555')
    for i, (txt, align) in enumerate([
        (milestone, WD_ALIGN_PARAGRAPH.LEFT),
        (stat,      WD_ALIGN_PARAGRAPH.CENTER),
        (notes,     WD_ALIGN_PARAGRAPH.LEFT),
    ]):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(txt)
        set_font(r, size=10,
                 bold=(i == 1),
                 color=tuple(int(color[j:j+2], 16) for j in (0,2,4)) if i == 1 else None)
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

table_caption('8.1', 'RTK Positioning module milestone status. ✔ Complete · ⏳ Pending · ○ Not started.')

# ═══════════════════════════════════════════════════════════════════════════
# 9. NEXT STEPS
# ═══════════════════════════════════════════════════════════════════════════
heading1('9.   Next Steps')

body('The following actions are planned for the next phase of the RTK Positioning module:')

bullet('Push all simulation results — Level 1, Level 2, and Level 3 data, analysis scripts, and generated figures — to the team GitHub repository to make the complete module history available to all team members.')
bullet('Execute Phase 3 RTKLIB validation: download real BeiDou (BDS) RINEX observation and navigation files from a pair of IGS MGEX reference stations separated by a known baseline distance. Run RTKLIB\'s rnx2rtkp post-processing engine configured for BeiDou-only positioning to obtain GNSS Only, RTK Float, and RTK Fixed accuracy statistics from real measurements. Compare these against the simulation noise model parameters used in Levels 1–3 to validate or recalibrate the simulation.')
bullet('Address the precision landing threshold gap identified in Level 3: the system achieved 31.5% within 0.3 m during the landing phase, well short of the 80% target. This can be approached either by extending the landing phase window (to allow re-convergence time) or by implementing a re-initialisation protocol that actively restores RTK Fixed lock upon entering the landing phase.')
bullet('Integrate the RTK positioning node output with the navigation and mission planning modules being developed by other team members, ensuring that RTK-corrected position data and the mission viability state signal are correctly consumed by the path planner within the shared simulation environment.')
bullet('Conduct a full system-level integration test within the Gazebo environment, running all team modules simultaneously, to verify that the RTK positioning output does not introduce timing or communication bottlenecks across the ROS 2 node graph.')
bullet('Investigate real hardware options: identify RTK-capable GNSS modules (e.g., u-blox F9P, Emlid Reach M2) for hardware-in-the-loop (HIL) testing to quantify the performance gap between simulation and physical measurements.')

spacer(20)

# ── Final save ────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f'Report saved: {OUT_PATH}')
